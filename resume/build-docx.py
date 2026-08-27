"""Build ATS-safe .docx and .txt resumes from resume_content.py.

Section order matches build-html.py exactly, so the three formats
present the same document.

ATS rules followed here:
  - Single column. No tables, text boxes, columns or shapes.
  - No document header/footer - many parsers skip that region entirely.
  - Contact details sit in the body, on one line, as real text.
  - Standard section headings as plain uppercase paragraphs, no
    letter-spacing (that makes glyphs extract individually).
  - Real bullet lists via the List Bullet style.
  - One standard font (Calibri) at 11pt body.
  - No images, icons or text rendered inside graphics.

Usage:  python resume/build-docx.py
"""

import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import resume_content as C  # noqa: E402

OUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "public")
BASENAME = "Dev-Shah-Senior-QA-Engineer-SDET"
INK = RGBColor(0x14, 0x14, 0x0F)
GREY = RGBColor(0x45, 0x44, 0x3D)
BODY = 11
SMALL = 10


def sp(p, before=0, after=4, line=1.12):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    return p


def build_docx():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(BODY)
    normal.font.color.rgb = INK

    for section in doc.sections:
        # 0.7in margins, matching the PDF.
        section.top_margin = section.bottom_margin = Pt(50)
        section.left_margin = section.right_margin = Pt(50)

    def heading(text):
        p = sp(doc.add_paragraph(), before=11, after=4)
        r = p.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(13)
        return p

    def para(text, size=BODY, bold_lead=None, colour=None, after=4):
        p = sp(doc.add_paragraph(), after=after)
        if bold_lead:
            r = p.add_run(bold_lead)
            r.bold = True
            r.font.size = Pt(size)
        r = p.add_run(text)
        r.font.size = Pt(size)
        if colour:
            r.font.color.rgb = colour
        return p

    def bullets(items):
        for it in items:
            bp = doc.add_paragraph(it, style="List Bullet")
            sp(bp, after=2)
            for run in bp.runs:
                run.font.size = Pt(BODY)

    # --- name / title / contact --------------------------------------
    p = sp(doc.add_paragraph(), after=2)
    r = p.add_run(C.NAME)
    r.bold = True
    r.font.size = Pt(20)

    p = sp(doc.add_paragraph(), after=2)
    r = p.add_run(C.TITLE)
    r.bold = True
    r.font.size = Pt(12)

    para("  |  ".join(C.CONTACT), size=SMALL, colour=GREY, after=2)

    # --- summary ------------------------------------------------------
    heading("Summary")
    para(C.SUMMARY)

    # --- core skills --------------------------------------------------
    heading("Core Skills")
    for label, items in C.SKILLS:
        para(items, bold_lead=f"{label}: ", after=3)

    # --- experience ---------------------------------------------------
    heading("Professional Experience")
    for job in C.EXPERIENCE:
        p = sp(doc.add_paragraph(), before=7, after=1)
        r = p.add_run(f"{job['role']} - {job['org']}")
        r.bold = True
        r.font.size = Pt(11.5)
        para(job["meta"], size=SMALL, colour=GREY, after=3)
        bullets(job["bullets"])

    # --- key projects -------------------------------------------------
    heading("Key Projects")
    for pr in C.PROJECTS:
        p = sp(doc.add_paragraph(), before=6, after=1)
        r = p.add_run(pr["name"])
        r.bold = True
        r.font.size = Pt(11.5)
        para(f"{pr['org']}  |  {pr['stack']}", size=SMALL, colour=GREY, after=2)
        if pr.get("compact"):
            p = sp(doc.add_paragraph(), after=3)
            p.add_run(pr["what"] + " ").font.size = Pt(BODY)
            r = p.add_run("Contribution: ")
            r.bold = True
            r.font.size = Pt(BODY)
            p.add_run(pr["did"]).font.size = Pt(BODY)
        else:
            para(pr["what"], after=2)
            para(pr["did"], bold_lead="Contribution: ", after=3)

    # --- education ----------------------------------------------------
    heading("Education, Training and Languages")
    para(C.EDUCATION, after=3)
    p = sp(doc.add_paragraph(), after=3)
    r = p.add_run("Training: "); r.bold = True; r.font.size = Pt(BODY)
    p.add_run(C.TRAINING + ". ").font.size = Pt(BODY)
    r = p.add_run("Languages: "); r.bold = True; r.font.size = Pt(BODY)
    p.add_run(C.LANGUAGES).font.size = Pt(BODY)

    for p in doc.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # never justify: ATS and readability

    path = os.path.join(OUT_DIR, BASENAME + ".docx")
    doc.save(path)
    return path


def build_txt():
    """Plain-text fallback - the format no parser can misread."""
    L = [C.NAME, C.TITLE, " | ".join(C.CONTACT), "", "SUMMARY", C.SUMMARY, "", "CORE SKILLS"]
    L += [f"{label}: {items}" for label, items in C.SKILLS]
    L += ["", "PROFESSIONAL EXPERIENCE"]
    for job in C.EXPERIENCE:
        L += ["", f"{job['role']} - {job['org']}", job["meta"]]
        L += [f"- {b}" for b in job["bullets"]]
    L += ["", "KEY PROJECTS"]
    for pr in C.PROJECTS:
        L += ["", pr["name"], f"{pr['org']} | {pr['stack']}", pr["what"],
              f"Contribution: {pr['did']}"]
    L += ["", "EDUCATION, TRAINING AND LANGUAGES", C.EDUCATION,
          f"Training: {C.TRAINING}", f"Languages: {C.LANGUAGES}"]

    path = os.path.join(OUT_DIR, BASENAME + ".txt")
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(L) + "\n")
    return path


if __name__ == "__main__":
    os.makedirs(OUT_DIR, exist_ok=True)
    print("wrote", build_docx())
    print("wrote", build_txt())
